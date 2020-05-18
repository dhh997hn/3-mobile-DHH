def information(phone_number,ip_number):
    import network_billing, call_billing
    netwk_invoice = network_billing.idata(ip_number)
    a,b = call_billing.idata(phone_number)
    call_invoice = a+b
    ssum = netwk_invoice + call_invoice
    tax = ssum * 0.18
    last_sum = ssum + tax
    netwk_invoice = '%.2f' % netwk_invoice
    call_invoice = '%.2f' % call_invoice
    ssum = '%.2f' % ssum
    tax = '%.2f' % tax
    last_sum = '%.2f' % last_sum
    data = {}
    data['Банк получателя'] = r'АО "Стоун банк" Г.МОСКВА'
    data['Получатель'] = r'ООО "Василек"'
    data['ИНН'] = 'ИНН 7722737766'
    data['КПП'] = 'КПП 772201001'
    data['БИК'] = '044525700'
    data['1Сч. №'] = '30101810200000000700'
    data['2Сч. №'] = '40702810900000002453'
    data['3Сч. №'] = '82'
    data['день'] = '12 '
    data['месяц'] = 'Апреля '
    data['год'] = '20'
    data['Поставщик'] = r'ООО "Василек",ИНН 7722737753,КПП 773301001,109052,Москва г., ДОБРЫНИНСКАЯ ул,дом № 70,корпус 2, тел.:'
    data['Заказчик'] = r'ООО "ЛАГУНА",ИНН 7714037378,КПП 777550001,119361,Москва г.,ТУЛЬСКАЯ ул.,дом № 4,строение 1'
    data['Основание'] = '№ 20022016 от 12.02.2020'
    data['№'] = ['1','2']
    data['товары'] = ['Тарификация услуг "Телефония"','Тарификация услуг "Интернет"']
    data['количество'] = ['','']
    data['ед.'] = ['','']
    data['цена'] = [call_invoice,netwk_invoice]
    data['сумма'] = [call_invoice,netwk_invoice]
    data['Итого'] = ssum
    data['НДС'] = tax
    data['Всего к оплату'] = last_sum
    data['Всего наименований'] = 'Всего наименований 2, на сумму '+ last_sum +' руб.'
    data['Сумма прописью'] = ''
    return data
def pdf(phone_number,ip_number):
    data = information(phone_number,ip_number)
    import comtypes.client
    wdFormatPDF = 17

    in_file = r'H:\study\nam 3\ki 2\Mobil\lab3\ha\form.docx'
    out_file = r'H:\study\nam 3\ki 2\Mobil\lab3\ha\invoice.docx'
    pdf_file = r'H:\study\nam 3\ki 2\Mobil\lab3\ha\invoice.pdf'

    from docx import Document
    doc = Document(in_file)

    #___Таблица 1__
    table = doc.tables[0]
    table.cell(0,0).paragraphs[0].text = data['Банк получателя']
    table.cell(0,3).paragraphs[0].text = data['БИК']
    table.cell(2,3).paragraphs[0].text = data['1Сч. №']
    table.cell(3,0).paragraphs[0].text = data['ИНН']
    table.cell(3,1).paragraphs[0].text = data['КПП']
    table.cell(3,3).paragraphs[0].text = data['2Сч. №']
    table.cell(4,0).paragraphs[0].text = data['Получатель']

    #___Таблица 2__
    table = doc.tables[1]
    inline = table.cell(0,0).paragraphs[0].runs
    text = inline[1].text.replace('___',data['3Сч. №'])
    inline[1].text = text
    text = inline[3].text.replace('___',data['день'])
    inline[3].text = text
    text = inline[5].text.replace('_______',data['месяц'])
    inline[5].text = text
    text = inline[7].text.replace('___',data['год'])
    inline[7].text = text

    #___Таблица 3__
    table = doc.tables[2]
    table.cell(0,1).paragraphs[0].text = data['Поставщик']
    table.cell(1,1).paragraphs[0].text = data['Заказчик']
    table.cell(2,1).paragraphs[0].text = data['Основание']

    #___Таблица 4__
    table = doc.tables[3]
    for i in range(len(data['№'])):
        if i>0: table.add_row()
        table.cell(i+1,0).text = data['№'][i]
        table.cell(i+1,1).text = data['товары'][i]
        table.cell(i+1,2).text = data['количество'][i]
        table.cell(i+1,3).text = data['ед.'][i]
        table.cell(i+1,4).text = data['цена'][i]
        table.cell(i+1,5).text = data['сумма'][i]

    #___Таблица 5__
    table = doc.tables[4]
    table.cell(0,1).paragraphs[0].text = data['Итого']
    table.cell(1,1).paragraphs[0].text = data['НДС']
    table.cell(2,1).paragraphs[0].text = data['Всего к оплату']
    table.cell(3,0).paragraphs[1].text = data['Всего наименований']
    table.cell(4,0).paragraphs[0].text = data['Сумма прописью']

   
    doc.save(out_file)

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(out_file)
    doc.SaveAs(pdf_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    
# variant 3
#pdf('915783624','192.168.250.27')
#create_invoice.pdf('915783624','192.168.250.27')